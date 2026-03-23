"""
Generate Coach M8 Competitive Analysis & Market Position document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Constants ──────────────────────────────────────────────────────────
BLUE_ACCENT = RGBColor(0x1A, 0x56, 0xDB)  # Coach M8 blue
DARK_BLUE = RGBColor(0x0D, 0x2B, 0x6B)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY_BG = "F2F2F2"
HEADER_BG = "1A56DB"
ALT_ROW_BG = "EBF0FA"
FONT_NAME = "Calibri"
OUTPUT_PATH = "/Users/bistrocloud/Documents/opsnerve-performance-analyzer/Coach_M8_Competitive_Analysis.docx"


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=Pt(9), alignment=None, font_name=FONT_NAME):
    """Set formatted text in a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_table_border(table):
    """Set thin borders on all table cells."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_heading_styled(doc, text, level=1):
    """Add a heading with custom blue styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else BLUE_ACCENT
        run.font.name = FONT_NAME
    return h


def add_paragraph_styled(doc, text, bold=False, italic=False, size=Pt(11), color=DARK_GRAY, space_after=Pt(6)):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.space_after = space_after
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_bold_body(doc, bold_text, normal_text, size=Pt(11)):
    """Add a paragraph with bold prefix and normal continuation."""
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r1 = p.add_run(bold_text)
    r1.font.name = FONT_NAME
    r1.font.size = size
    r1.font.bold = True
    r1.font.color.rgb = DARK_GRAY
    r2 = p.add_run(normal_text)
    r2.font.name = FONT_NAME
    r2.font.size = size
    r2.font.color.rgb = DARK_GRAY
    return p


def add_numbered_list(doc, items, bold_prefix=False):
    """Add numbered list items."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.style = doc.styles['List Number']
        p.space_after = Pt(3)
        if bold_prefix and ": " in item:
            parts = item.split(": ", 1)
            r1 = p.add_run(parts[0] + ": ")
            r1.font.name = FONT_NAME
            r1.font.size = Pt(11)
            r1.font.bold = True
            r1.font.color.rgb = DARK_GRAY
            r2 = p.add_run(parts[1])
            r2.font.name = FONT_NAME
            r2.font.size = Pt(11)
            r2.font.color.rgb = DARK_GRAY
        else:
            run = p.add_run(item)
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY


def create_cover_page(doc):
    """Create a professional cover page."""
    # Add spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.space_after = Pt(0)

    # Title line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(0)
    run = p.add_run("COACH M8")
    run.font.name = FONT_NAME
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = BLUE_ACCENT

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(8)
    run = p.add_run("AI-Powered Football Performance Intelligence")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.color.rgb = MEDIUM_GRAY

    # Horizontal rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(8)
    run = p.add_run("_" * 60)
    run.font.color.rgb = BLUE_ACCENT
    run.font.size = Pt(11)

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(30)
    run = p.add_run("Competitive Analysis & Market Position")
    run.font.name = FONT_NAME
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    run = p.add_run(datetime.datetime.now().strftime("%B %Y"))
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_GRAY

    # Prepared by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    run = p.add_run("Prepared by")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = MEDIUM_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    run = p.add_run("Opsnerve Technologies")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("In partnership with Ahmed Hossam (Mido)")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.italic = True

    # Confidential notice at bottom
    for _ in range(4):
        doc.add_paragraph().space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()


def create_toc(doc):
    """Create a Table of Contents placeholder."""
    add_heading_styled(doc, "Table of Contents", level=1)

    toc_items = [
        ("1.", "Executive Summary", "3"),
        ("2.", "Competitive Landscape", "4"),
        ("3.", "Workflow Comparison", "6"),
        ("4.", "Price Analysis", "8"),
        ("5.", "The Mido Factor -- Strategic Advantage", "9"),
        ("6.", "Risk Assessment", "11"),
        ("7.", "Conclusion", "12"),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.space_after = Pt(6)
        # Section number
        r1 = p.add_run(f"{num}  ")
        r1.font.name = FONT_NAME
        r1.font.size = Pt(12)
        r1.font.bold = True
        r1.font.color.rgb = BLUE_ACCENT
        # Title
        r2 = p.add_run(title)
        r2.font.name = FONT_NAME
        r2.font.size = Pt(12)
        r2.font.color.rgb = DARK_GRAY

    doc.add_page_break()


def create_competitive_table(doc):
    """Create the large competitive landscape comparison table."""
    # Data: each row is [Feature, Coach M8, Catapult, STATSports, Kinexon, Playermaker, Zone7, Whoop]
    rows_data = [
        ["Heart Rate Monitoring", "BLE chest straps + ESP32", "GPS vest with HR", "GPS vest with HR", "UWB + HR", "No HR", "No hardware", "Wrist band"],
        ["GPS/Position Tracking", "CV pipeline from video", "10Hz GPS", "18Hz GPS", "cm-level UWB", "Boot sensor only", "No hardware", "No position"],
        ["Distance/Speed/Sprints", "From video (sub-meter)", "From GPS (3-5m error)", "From GPS (3-5m error)", "From UWB", "Step-based only", "Data integration", "No"],
        ["Video Analysis (CV)", "Own YOLO pipeline", "No", "No", "No", "No", "No", "No"],
        ["Tactical Analysis", "Formations, pressing, transitions", "No", "No", "Basic", "No", "No", "No"],
        ["AI Reports & Insights", "Claude AI -- NL queries, reports", "Raw data only", "Raw data only", "Basic", "No", "Injury prediction only", "No"],
        ["AI Coaching Chat", "Ask anything, full data context", "No", "No", "No", "No", "No", "No"],
        ["Squad Builder", "AI-recommended starting XI", "No", "No", "No", "No", "No", "No"],
        ["Weekly Planner", "AI-generated training schedule", "No", "No", "No", "No", "No", "No"],
        ["Injury Risk (ACWR)", "Auto-calculated + AI alerts", "Manual calculation", "Manual calculation", "Basic", "No", "ML prediction", "Basic"],
        ["Live HR Dashboard", "Real-time during sessions", "With their hardware", "With their hardware", "With their hardware", "No", "No", "No"],
        ["WhatsApp Delivery", "Alerts, reports, NL queries", "No", "No", "No", "No", "No", "No"],
        ["Youth Development", "Age-appropriate benchmarks", "Adult metrics only", "Adult metrics only", "Adult metrics only", "No", "No", "No"],
        ["Camera Agnostic", "Works with any video", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A"],
        ["Hardware Cost", "~$1,500", "$100K+ in sub", "$50-80K in sub", "$150-300K venue", "$200/player", "No hardware", "$30/player/mo"],
        ["Monthly Cost", "$1,000/mo", "~$8,000+/mo", "~$4-7K/mo", "Custom enterprise", "~$200/player/yr", "~$2.5-8K/mo", "$30/player/mo"],
        ["Minimum Contract", "None", "2-3 years", "1-2 years", "Multi-year", "Annual", "Annual", "Monthly"],
        ["Target Market", "Academies ($5-25K)", "Pro clubs ($100K+)", "Pro clubs", "Elite clubs/leagues", "Individual athletes", "Pro clubs", "Individual athletes"],
    ]

    headers = [
        "Feature",
        "Coach M8\n($25K + $1K/mo)",
        "Catapult Vector\n($100K+/yr)",
        "STATSports Apex\n($50-80K/yr)",
        "Kinexon\n($150-300K+)",
        "Playermaker\n($200+/yr/player)",
        "Zone7 AI\n($30-100K/yr)",
        "Whoop\n($30/player/mo)",
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_border(table)

    # Header row
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=WHITE, size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    yes_keywords = ["BLE", "CV pipeline", "From video", "Own YOLO", "Formations", "Claude AI", "Ask anything",
                    "AI-recommended", "AI-generated", "Auto-calculated", "Real-time", "Alerts", "Age-appropriate",
                    "Works with", "~$1,500", "$1,000/mo", "None", "Academies"]
    no_keywords = ["No HR", "No hardware", "No position", "No", "Boot sensor only", "Step-based only", "N/A"]

    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            # Alternating row background
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

            if j == 0:
                # Feature name column
                set_cell_text(cell, val, bold=True, size=Pt(8), color=DARK_BLUE)
            elif j == 1:
                # Coach M8 column - highlight green-ish
                prefix = ""
                if val not in ["~$1,500", "$1,000/mo", "None", "Academies ($5-25K)"]:
                    prefix = ""
                set_cell_text(cell, val, size=Pt(8), color=RGBColor(0x0D, 0x6B, 0x2B))
                cell.paragraphs[0].runs[0].font.bold = True
            else:
                # Competitor columns
                is_no = val in ["No", "No HR", "No hardware", "No position", "N/A", "Boot sensor only", "Step-based only", "Data integration"]
                text_color = RGBColor(0xCC, 0x00, 0x00) if is_no else DARK_GRAY
                set_cell_text(cell, val, size=Pt(8), color=text_color)

    # Set column widths
    widths = [Inches(1.3), Inches(1.4), Inches(1.1), Inches(1.1), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]


def create_price_table(doc):
    """Create the price analysis comparison table."""
    headers = ["Metric", "Coach M8", "Catapult", "STATSports", "Zone7"]
    rows_data = [
        ["Year 1 Total", "$37,000", "$100,000+", "$50-80,000", "$30-100,000"],
        ["Year 2 Total", "$12,000", "$100,000+", "$50-80,000", "$30-100,000"],
        ["Year 3 Total", "$12,000", "$100,000+", "$50-80,000", "$30-100,000"],
        ["3-Year TCO", "$61,000", "$300,000+", "$150-240,000", "$90-300,000"],
        ["Staff Required", "0 (AI replaces analyst)", "1 full-time analyst", "1 full-time analyst", "1 data scientist"],
        ["Analyst Salary (3yr)", "$0", "$72-144,000", "$72-144,000", "$144,000+"],
        ["Real 3-Year Cost", "$61,000", "$372-444,000", "$222-384,000", "$234-444,000"],
        ["Cost per Player/Year\n(65 players)", "$189", "$1,538-2,282", "$769-1,231", "$462-1,538"],
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    # Header
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            if j == 0:
                set_cell_text(cell, val, bold=True, size=Pt(10), color=DARK_BLUE)
            elif j == 1:
                # Coach M8 column - green highlight
                set_cell_text(cell, val, bold=True, size=Pt(10), color=RGBColor(0x0D, 0x6B, 0x2B))
            else:
                set_cell_text(cell, val, size=Pt(10), color=DARK_GRAY)

    widths = [Inches(1.8), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]


def create_risk_table(doc):
    """Create the risk assessment table."""
    headers = ["Risk", "Mitigation"]
    rows_data = [
        ["CV pipeline accuracy vs GPS", "Start with HR-only (proven), add CV progressively. HR alone delivers 80% of value."],
        ["Academy budget constraints", "$25K is within reach for serious academies. ROI: replaces $24-48K/yr analyst salary."],
        ["Catapult enters AI space", "They're hardware companies -- AI is not their DNA. Our AI advantage grows with data."],
        ["Data privacy (youth players)", "Parental consent mechanism built in. Compliant with Egypt's data protection law."],
        ["Mido's time commitment", "Mido is the door opener, not the operator. Tech team runs the product."],
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            if j == 0:
                set_cell_text(cell, val, bold=True, size=Pt(10), color=RGBColor(0xCC, 0x00, 0x00))
            else:
                set_cell_text(cell, val, size=Pt(10), color=DARK_GRAY)

    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(5.0)


def add_footer(doc):
    """Add footer with confidential notice and page numbers."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(4)

        # Confidential text
        run = p.add_run("Confidential -- Coach M8 / Opsnerve Technologies")
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = MEDIUM_GRAY
        run.font.italic = True

        # Page number
        run2 = p.add_run("    |    Page ")
        run2.font.name = FONT_NAME
        run2.font.size = Pt(8)
        run2.font.color.rgb = MEDIUM_GRAY

        # Add page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run3 = p.add_run()
        run3._r.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run4 = p.add_run()
        run4._r.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5 = p.add_run()
        run5._r.append(fldChar2)


def add_header_bar(doc):
    """Add a header with Coach M8 branding."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = p.add_run("Coach M8")
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = BLUE_ACCENT

        run2 = p.add_run("  |  Competitive Analysis")
        run2.font.name = FONT_NAME
        run2.font.size = Pt(9)
        run2.font.color.rgb = MEDIUM_GRAY

        # Bottom border on header paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A56DB"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)


# ── Main Document Generation ──────────────────────────────────────────
def generate():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── Cover Page ──
    create_cover_page(doc)

    # ── Table of Contents ──
    create_toc(doc)

    # ── Section 1: Executive Summary ──
    add_heading_styled(doc, "1. Executive Summary", level=1)

    add_paragraph_styled(doc,
        "Coach M8 is an AI-powered football performance analysis and squad management platform "
        "built for youth football academies. Co-founded with Ahmed Hossam (Mido), former Egyptian "
        "international and co-owner of The Maker Football Incubator. Coach M8 combines wearable "
        "heart rate monitoring, computer vision video analysis, AI coaching intelligence, and squad "
        "management into a single platform -- at 1/10th the cost of professional alternatives."
    )

    add_bold_body(doc, "Pricing: ", "$25,000 one-time setup + $1,000/month subscription.")

    doc.add_page_break()

    # ── Section 2: Competitive Landscape ──
    add_heading_styled(doc, "2. Competitive Landscape", level=1)
    add_paragraph_styled(doc,
        "The table below compares Coach M8 against every major competitor across features, pricing, "
        "and target market. Green values indicate Coach M8 advantages; red values indicate competitor gaps.",
        size=Pt(10), italic=True, color=MEDIUM_GRAY
    )

    # Switch to landscape for the wide table
    new_section = doc.add_section(WD_ORIENT.LANDSCAPE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Inches(11)
    new_section.page_height = Inches(8.5)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)
    new_section.top_margin = Cm(2.0)
    new_section.bottom_margin = Cm(2.0)

    create_competitive_table(doc)

    # Back to portrait
    new_section2 = doc.add_section(WD_ORIENT.PORTRAIT)
    new_section2.orientation = WD_ORIENT.PORTRAIT
    new_section2.page_width = Inches(8.5)
    new_section2.page_height = Inches(11)
    new_section2.left_margin = Cm(2.0)
    new_section2.right_margin = Cm(2.0)
    new_section2.top_margin = Cm(2.0)
    new_section2.bottom_margin = Cm(2.0)

    # ── Section 3: Workflow Comparison ──
    add_heading_styled(doc, "3. Workflow Comparison", level=1)

    add_heading_styled(doc, "Catapult/STATSports Workflow (Current Industry Standard)", level=2)
    add_numbered_list(doc, [
        "Coach runs training -- players wear GPS vests",
        "Data syncs post-session via proprietary docking station",
        "Performance analyst opens Catapult/STATSports software",
        "Analyst manually creates reports (30-60 min per session)",
        "Analyst emails reports to coaching staff",
        "Coach reads reports next day",
        "No AI insights, no squad recommendations, no training planning",
    ])

    add_heading_styled(doc, "Coach M8 Workflow (AI-First)", level=2)
    add_numbered_list(doc, [
        "Coach opens Coach M8 dashboard -- AI Daily Briefing shows today's plan, who needs rest, team readiness",
        "Players wear BLE chest straps -- data streams live to coach's tablet during session",
        "Post-session: AI auto-generates reports in 30 seconds (not 60 minutes)",
        'Coach gets WhatsApp alert: "Session complete. 2 players flagged for load management."',
        'Coach clicks "AI Report" -- full analysis with per-player recommendations',
        "Squad Builder recommends starting XI for Saturday with data-backed reasoning",
        "Weekly Planner generates next week's training schedule based on load data",
        'Coach chats with AI: "Should Ahmed play the full match Saturday?" -- instant answer with data',
    ])

    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run("Time savings: 60+ min/session --> 30 seconds. No dedicated analyst needed.")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = BLUE_ACCENT

    doc.add_page_break()

    # ── Section 4: Price Analysis ──
    add_heading_styled(doc, "4. Price Analysis", level=1)
    create_price_table(doc)

    doc.add_paragraph()  # spacer
    add_paragraph_styled(doc,
        "Coach M8 is 5-7x cheaper than the nearest competitor AND includes features they don't have "
        "(AI coaching, squad management, weekly planning).",
        bold=True, size=Pt(12), color=BLUE_ACCENT
    )

    doc.add_page_break()

    # ── Section 5: The Mido Factor ──
    add_heading_styled(doc, "5. The Mido Factor -- Strategic Advantage", level=1)
    add_heading_styled(doc, "Ahmed Hossam (Mido) as Co-Founder -- What This Means", level=2)

    mido_points = [
        ("Credibility in Egyptian Football: ",
         "Mido is one of Egypt's most recognized footballers. His name on Coach M8 gives instant credibility "
         "with every academy director in Egypt and the Middle East."),
        ("The Maker as Proving Ground: ",
         "80 players, 65 staff, 4 locations. If Coach M8 works at The Maker, Mido can demonstrate results to "
         'any academy. "We used this on our own players -- here are the results."'),
        ("Network Access: ",
         "Mido's relationships with Egyptian FA, Saudi Pro League clubs, UAE academies, and European scouts "
         "create distribution channels no startup could buy."),
    ]
    for bold_t, normal_t in mido_points:
        add_bold_body(doc, bold_t, normal_t)

    # Go-to-Market
    add_heading_styled(doc, "Go-to-Market Strategy", level=2)
    phases = [
        "Phase 1: Prove at The Maker (done -- demo shows 1-month improvement)",
        "Phase 2: 5 Egyptian academies (Mido's network)",
        "Phase 3: Gulf region (UAE, Saudi, Qatar -- high spending, English-speaking)",
        "Phase 4: European academy partnerships (via scouting networks)",
    ]
    for phase in phases:
        add_bold_body(doc, phase.split(":")[0] + ": ", phase.split(": ", 1)[1])

    # Revenue Projections
    add_heading_styled(doc, "Revenue Projections", level=2)
    projections = [
        "Year 1: 5 academies x $37K = $185K",
        "Year 2: 20 academies x $12K (recurring) + 15 new x $37K = $795K",
        "Year 3: 50 academies x $12K + 30 new x $37K = $1.71M",
        "With wearable hardware (future): $3-5M ARR by Year 3",
    ]
    for proj in projections:
        p = doc.add_paragraph(proj, style='List Bullet')
        p.space_after = Pt(3)
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY

    # Competitive Moat
    add_heading_styled(doc, "Competitive Moat", level=2)
    moat_items = [
        "Catapult/STATSports sell hardware. Coach M8 sells intelligence.",
        "They charge $100K for data. We charge $25K for answers.",
        "Their customers need analysts. Our customers need a phone.",
        "Mido opens doors that cold calls never will.",
    ]
    for item in moat_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.space_after = Pt(3)
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE

    doc.add_page_break()

    # ── Section 6: Risk Assessment ──
    add_heading_styled(doc, "6. Risk Assessment", level=1)
    create_risk_table(doc)

    doc.add_page_break()

    # ── Section 7: Conclusion ──
    add_heading_styled(doc, "7. Conclusion", level=1)

    add_paragraph_styled(doc,
        "Coach M8 is not competing with Catapult on hardware. It's competing on intelligence -- and winning. "
        "No existing product combines AI coaching, squad management, live monitoring, video analysis, and "
        "WhatsApp delivery at academy-affordable pricing. With Mido as co-founder, the go-to-market strategy "
        'is not "cold call 1000 academies" -- it\'s "Mido calls 5 friends who run academies."'
    )

    # Closing statement - highlighted
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "The question isn't whether Coach M8 can compete with Catapult.\n"
        "It's whether Catapult can compete with AI + Mido + academy pricing."
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BLUE_ACCENT

    # ── Headers and Footers ──
    add_header_bar(doc)
    add_footer(doc)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate()
